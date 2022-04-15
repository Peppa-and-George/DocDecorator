"""
获取文档中的内容元素：
    包含文档中的 paragraph,table,row,column,cell,run
"""
import docx.document
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.table import Table, _Row, _Column, _Cell


class DOCXElement(object):
    def __init__(self, file: str | docx.document.Document):
        """
        :param file: 需要获取元素的文档
        """
        if isinstance(file, str):
            self.doc = docx.Document(file)
        else:
            self.doc = file

        self.paragraphs = None
        self.tables = None
        self.runs = None

    def get_paragraphs_of_document(self) -> list[Paragraph]:
        """
        获取文档中的所有段落
        :return: 文档中的所有段落对象
        """
        self.paragraphs = self.doc.paragraphs
        return self.paragraphs

    def get_tables_of_document(self) -> list[Table]:
        """
        获取文档中的所有表格
        :return: 文档中的所有表格对象
        """
        return self.doc.tables

    def get_all_paragraphs(self) -> list[Paragraph]:
        """
        获取文章中所有的段落对象
        :return: 段落对象列表
        """
        paragraphs = []
        paragraphs.extend(self.get_paragraphs_of_document())
        for table in self.get_tables_of_document():
            paragraphs.extend(self.get_paragraphs_of_table(table))
        return paragraphs

    def get_all_runs(self) -> list[Run]:
        """
        获取文本中所有的run对象
        :return: run对象列表
        """
        runs = []
        for paragraph in self.get_all_paragraphs():
            runs.extend(paragraph.runs)
        return runs

    def get_runs_of_paragraphs(self) -> list[Run]:
        """
        获取文档中所有段落中的run
        :return: 文档中所有段落的所有run对象
        """
        runs = []
        if self.paragraphs is None:
            self.get_paragraphs_of_document()
        for paragraph in self.paragraphs:
            runs.extend(paragraph.runs)
        return runs

    def get_runs_of_paragraph(self, paragraph: int | Paragraph) -> list[Run]:
        """
        获取指定段落中的run
        :param paragraph: 需要获取run的段落对象或者段落下标
        :return: 指定段落中的所有run
        """
        if isinstance(paragraph, int):
            if self.paragraphs is None:
                self.paragraphs = self.get_paragraphs_of_document()
            assert 0 <= paragraph < len(self.paragraphs), ValueError('段落下标不合规')
            paragraph = self.paragraphs[paragraph]
        else:
            paragraph = paragraph
        return paragraph.runs

    def get_paragraph(self, paragraph: Paragraph | int) -> Paragraph:
        """
        获取指定段落对象
        :param paragraph: 段落对象或者段落下标
        :return: 指定对罗对象
        """
        if isinstance(paragraph, int):
            return self.get_paragraph_by_index(paragraph)
        else:
            return paragraph

    def get_paragraph_by_index(self, index: int) -> Paragraph:
        """
        根据下标获取文档中的指定段落
        :param index: 段落下标
        :return: 下标为index的段对象
        """
        if self.paragraphs is None:
            self.get_paragraphs_of_document()
        assert 0 <= index < len(self.paragraphs), ValueError('段落下标不合规')
        return self.paragraphs[index]

    def get_table(self, table: Table | int) -> Table:
        """
        获取表格对象
        :param table: 表格对象或者表格下标
        :return: 指定表格对象
        """
        if isinstance(table, int):
            return self.get_table_by_index(table)
        else:
            return table

    def get_table_by_index(self, index: int) -> Table:
        """
        通过下标获取文档中的表格
        :param index: 表格下标
        :return: 下标为index的表格对象
        """
        if self.tables is None:
            self.get_tables_of_document()
        assert 0 <= index < len(self.tables), ValueError('表格下标不合格')
        return self.tables[index]

    @staticmethod
    def get_rows_of_table(table: Table):
        """
        获取指定表格中的所有的row
        :param table: 表格对象
        :return: 指定表格中的所有行
        """
        return table.rows

    @staticmethod
    def get_columns_of_table(table: Table):
        """
        获取指定表格中的所有的columns
        :param table: 表格对象
        :return: 指定表格中的所有列
        """
        return table.columns

    @classmethod
    def get_row_of_table_by_index(cls, table: Table, row_index: int) -> _Row:
        """
        通过下标获取表格中的行
        :param table: 表格对象
        :param row_index: 行下标
        :return: 指定行
        """
        return cls.get_rows_of_table(table)[row_index]

    def get_column_of_table_by_index(self, table: Table, column_index: int) -> _Column:
        """
        通过下标获取表格中的行
        :param table: 表格对象
        :param column_index: 列下标
        :return: 指定列
        """
        return self.get_columns_of_table(table)[column_index]

    @staticmethod
    def get_cells_of_table(table: Table) -> list[_Cell]:
        """
        获取指定表格的所有单元格
        :param table: 表格对象
        :return: 指定表格的所有单元格
        """
        return table._cells

    @staticmethod
    def get_cells_of_row(row: _Row) -> list[_Cell]:
        """
        获取指定行里的所有单元格
        :param row: 指定行
        :return: 单元格列表
        """
        return row.cells

    @staticmethod
    def get_cells_of_column(column: _Column) -> list[_Cell]:
        """
        获取指定列的单元格
        :param column: 指定列对象
        :return: 单元格对象列表
        """
        return column.cells

    @staticmethod
    def get_paragraphs_of_cell(cell: _Cell) -> list[Paragraph]:
        """
        获取致电给单元格里的所有段落
        :param cell: 指定单元格
        :return: 段落对象列表
        """
        return cell.paragraphs

    @classmethod
    def get_runs_of_cell(cls, cell: _Cell) -> list[Run]:
        """
        获取指定单元格里的所有run
        :param cell: 指定单元格对象
        :return: run列表
        """
        runs = []
        for paragraph in cls.get_paragraphs_of_cell(cell):
            runs.extend(paragraph.runs)
        return runs

    @classmethod
    def get_paragraphs_of_table(cls, table: Table) -> list[Paragraph]:
        """
        获取指定表格的所有段落
        :param table: 表格对象
        :return: 段落列表
        """
        paragraphs = []
        for cell in cls.get_cells_of_table(table):
            paragraphs.extend(cls.get_paragraphs_of_cell(cell))
        return paragraphs

    @classmethod
    def get_runs_of_table(cls, table: Table) -> list[Run]:
        """
        获取指定表格的多有run对象
        :param table: 表格对象
        :return: run对象列表
        """
        runs = []
        for paragraph in cls.get_paragraphs_of_table(table):
            runs.extend(paragraph.runs)
        return runs

    @classmethod
    def get_paragraphs_of_row(cls, row: _Row):
        """
        获取指定行的所有段落
        :param row: 行对象
        :return: 单元格对象列表
        """
        paragraphs = []
        for cell in cls.get_cells_of_row(row):
            paragraphs.extend(cell.paragraphs)
        return paragraphs

    @classmethod
    def get_paragraphs_of_column(cls, column: _Column):
        """
        获取指定列的所有段落
        :param column: 列对象
        :return: 单元格对象列表
        """
        paragraphs = []
        for cell in cls.get_cells_of_column(column):
            paragraphs.extend(cell.paragraphs)
        return paragraphs

    def get_runs_of_column(self, column: _Column) -> list[Run]:
        """
        获取指定列的多有run对象
        :param column: 列对象
        :return: run对象列表
        """
        runs = []
        for paragraph in self.get_paragraphs_of_column(column):
            runs.extend(paragraph.runs)
        return runs

    def get_runs_of_row(self, row: _Row) -> list[Run]:
        """
        获取指定列的多有run对象
        :param row: 行对象
        :return: run对象列表
        """
        runs = []
        for paragraph in self.get_paragraphs_of_row(row):
            runs.extend(paragraph.runs)
        return runs

    @classmethod
    def get_cells_of_table_element(cls, element: _Row | _Column | Table | _Cell) -> list[_Cell]:
        """
        获取表格元素中的所有单元格
        :param element: 表格元素对象
        :return: 单元格列表
        """
        if isinstance(element, Table):
            return element._cells()
        elif isinstance(element, _Cell):
            return [element]
        else:
            return element.cells

    @classmethod
    def get_paragraphs_of_table_element(cls, element: _Row | _Column | Table | _Cell) -> list[Paragraph]:
        """
        获取表格元素中的所有单元格
        :param element: 表格元素对象
        :return: 段落对象列表
        """
        paragraphs = []
        for cell in cls.get_cells_of_table_element(element):
            paragraphs.extend(cell.paragraphs)
        return paragraphs

    @classmethod
    def get_runs_in_table_element(cls, element: _Row | _Column | Table | _Cell) -> list[Run]:
        """
        获取表格元素中的所有单元格
        :param element: 表格元素对象
        :return: run对象列表
        """
        runs = []
        for paragraph in cls.get_paragraphs_of_table_element(element):
            runs.extend(paragraph.runs)
        return runs
