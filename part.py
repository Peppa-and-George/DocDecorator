"""
获取文档元素模块

- 获取文档的各个部分，如获取文档中所有表格，段落，run等
"""
import docx
from docx.table import Table, _Cell, _Row
from docx.text import run
from docx.text.paragraph import Paragraph
from docx.document import Document
from docx.text.run import Run


def open_document(doc: str) -> docx.Document:
    """
    打开一个word文档并返回
    :param doc: 文档路径
    :return: 文档对象
    """
    try:
        return docx.Document(doc)
    except IOError:
        raise IOError(f'打开文档{doc}失败')


def get_paragraphs(doc: Document | str) -> list[Paragraph]:
    """
    返回文档中的段落
    :param doc: 文档对象或者文档路径
    :return: 文档中的段落列表
    """
    if isinstance(doc, str):
        doc = open_document(doc)

    return doc.paragraphs


def get_tables(doc: Document | str) -> list[Table]:
    """
    返回文中所有表格
    :param doc: 文端对象或者文档路径
    :return:
    """
    if isinstance(doc, str):
        doc = open_document(doc)

    return doc.tables


def get_cells_in_tables(doc: Document | str) -> list[_Cell]:
    """
    获取文档里的所有单元格
    :param doc: 文档对象或者文档路径
    :return: 单元格列表
    """
    cells = []
    tables = get_tables(doc)

    for table in tables:
        cells.extend(get_cells_in_table(table))
    return cells


def get_cells_in_table(table: Table) -> list[_Cell]:
    """
    获取表格中的所有单元格

    :param table docx.table.Table
        需要获取单元的表格
    :return list[docx.table._Cell]
        单元个对象列表
    """
    cells = []
    for i in range(get_table_row_length(table)):
        cells.extend(table.row_cells(i))
    return cells


def get_table_row_length(table: Table) -> int:
    """
    获取表格的行数

    :param table docx.table.Table
        需要获取行数的表格对象
    :return int
        表格行数
    """
    return len(table.rows)


def get_table_column_length(table: Table) -> int:
    """
    获取表格的列数

    :param table docx.table.Table
        需要获取列数的表格对象
    :return int
        表格列数
    """
    return len(table.columns)


def get_paragraphs_in_tables(doc: Document | str) -> \
        list[docx.text.paragraph.Paragraph]:
    """
    获取文档表格中的所有段落
    :param doc: 文档或者文档路径
    :return: 表格中的段落列表
    """
    paragraphs = []
    for cell in get_cells_in_tables(doc):
        paragraphs.extend(cell.paragraphs)
    return paragraphs


def get_runs_in_paragraphs(doc: Document | str) -> list[run.Run]:
    """
    获取段落里的run
    :param doc: 文档对象或者文档路径
    :return: 文档段落中的所有run列表
    """
    runs = []
    for paragraph in get_paragraphs(doc):
        runs.extend(paragraph.runs)

    return runs


def get_runs_in_tables(doc: Document | str) -> list[run.Run]:
    """
    获取段落里的run
    :param doc: 文档对象或者文档路径
    :return: 表格中所有run列表
    """
    runs = []
    for paragraph in get_paragraphs_in_tables(doc):
        runs.extend(paragraph.runs)

    return runs


def get_all_paragraphs(doc: Document | str) -> list[Paragraph]:
    """
    获取文档中所有的段落
    :param doc: 文档对象或者文档路径
    :return: 文档中所有r段落列表
    """
    paragraphs = []
    paragraphs.extend(get_paragraphs(doc))
    paragraphs.extend(get_paragraphs_in_tables(doc))
    return paragraphs


def get_all_runs(doc: Document | str) -> list[run.Run]:
    """
    获取文档中所有的run
    :param doc: 文档对象或者文档路径
    :return: 文档中所有run列表
    """
    cells = []
    for paragraph in get_all_paragraphs(doc):
        cells.extend(paragraph.runs)
    return cells


def get_run_in_row(row: _Row) -> list[Run]:
    """
    获取行里面的单元格

    :param row docx.table._Row
        需要获取run的表格行
    """
    runs = []
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            runs.extend(paragraph.runs)
    return runs


def get_runs_in_cell(cell: _Cell) -> list[Run]:
    """
    获取单元格里面的run
    :param cell docx.table._Cell
        需要获取runs的单元格
    """
    runs = []
    for paragraph in cell.paragraphs:
        runs.extend(paragraph.runs)
    return runs


def get_row_in_table(table: Table, row_index: int) -> _Row:
    """
    获取表格中的某一行

    :param table docx.table.Table
        需要获取行的表格
    :param row_index int
        行的下标(索引)
    :return docx.table._Row
        行对象
    """
    assert 0 < row_index < get_table_row_length(table), \
        ValueError('row_index out of range')
    return table.rows[row_index]


def get_column_in_table(table: Table, column_index: int) -> _Row:
    """
    获取表格中的某一行

    :param table docx.table.Table
        需要获取行的表格
    :param column_index int
        列的下标(索引)
    :return docx.table._Row
        列对象
    """
    assert 0 < column_index < get_table_column_length(table), \
        ValueError('column_index out of range')
    return table.columns[column_index]


def get_runs_in_element(
        element: Document | Table | _Row | _Cell | Paragraph | Run | list[Run]
) -> list[Run]:
    """
    获取元素里面的所有run
    :param element
         元素
    """
    runs = []
    if isinstance(element, Document):
        runs = get_all_runs(element)
    elif isinstance(element, Table):
        for cell in get_cells_in_table(element):
            runs.extend(get_runs_in_cell(cell))
    elif isinstance(element, _Cell):
        runs = get_runs_in_cell(element)
    elif isinstance(element, Paragraph):
        runs = element.runs
    elif isinstance(element, _Row):
        for cell in element.cells:
            runs.extend(get_runs_in_cell(cell))
    elif isinstance(element, Run):
        runs.append(element)
    else:
        runs = element
    return runs


def get_paragraphs_in_table(table: Table) -> list[Paragraph]:
    """
    获取表格中所有的堕落

    :param table docx.table.Table
        需要获取段落的表格
    """
    paragraphs = []
    for cell in get_cells_in_table(table):
        paragraphs.extend(cell.paragraphs)
    return paragraphs
