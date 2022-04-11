# @Time: 16/3/2022 下午2:54
# @Author: kang

"""
修饰doc的模块

- 修饰文档内容
- 修改文本字体字号
- 蟹盖段落样式
"""

from docx.document import Document
import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.table import Table, _Row, _Cell, _Column
from docx.text.paragraph import Paragraph
from docx.text.run import Run
import part


class DocDecorate(object):
    """
    修饰一个docx文档，可以实现：
        1. 文本替换
        2. 添加图片
        3. 添加一个段落
        4. 添加表格行
        5. 设置元素的字体
        6. 设置元素字号
        7. 设置段落缩进
    """

    # 段落对齐方式
    P_CENTER = WD_PARAGRAPH_ALIGNMENT.CENTER
    P_LEFT = WD_PARAGRAPH_ALIGNMENT.LEFT
    P_RIGHT = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # 表格对齐方式
    T_CENTER = WD_TABLE_ALIGNMENT.CENTER
    T_LEFT = WD_TABLE_ALIGNMENT.LEFT
    T_RIGHT = WD_TABLE_ALIGNMENT.RIGHT

    def __init__(self, doc: Document | str):
        """
        :param docx.document.Document|str doc:
            需要修饰的文档，可传入Document对象，也可传入doc文档虽在路径
        """
        if isinstance(doc, str):
            self.doc = docx.Document(doc)
        else:
            self.doc = doc

        self.runs = part.get_all_runs(self.doc)
        self.paragraphs = part.get_paragraphs(self.doc)
        self.tables = part.get_tables(self.doc)
        self.cells = part.get_cells_in_tables(self.doc)

    def replace(self, context: dict):
        """
        替换文档中的文本

        :param  context dict
            需要替换的内容，例如：
            {
                old: new,
                ...
            }
        """
        # 获取文档中所有的段落中的run对象然后替换内容
        for run in self.runs:
            for key in context:
                if key in run.text:
                    run.text = run.text.replace(key, str(context.get(key)))

    def add_picture(self, context: dict):
        """
        在文档中特定文字处插入图片

        :param context dict
            需要替换的位置何替换图片的对应关系，例子：
            {
                text: pic_path,
                ...
            }
        """
        for run in self.runs:
            for key in context:
                if key in run.text:
                    try:
                        run.text = ''
                        run.add_picture(context.get(key), width=Cm(15))
                    except Exception:
                        raise Exception('添加图片失败')

    def append_row(self, table_index: int, row_data: list[str]) -> _Row:
        """
        在表格table上添加一个新的行，然后将row_data的数据填充到table中

        :param table_index int
            需要添加行的表格的下标
        :param row_data list[str]
            行里面的数据，数据为一个列表，代表行里面每个单元格的数据
        """

        try:
            row = self.tables[table_index].add_row()
            cells = row.cells
            assert len(cells) == len(row_data), ValueError('传入的数据与单元格数量不一致')
            for index in range(len(cells)):
                run = cells[index].paragraphs[0].add_run()
                run.text = str(row_data[index])
            self.tables = part.get_tables(self.doc)
            return row
        except Exception:
            raise Exception('添加行失败')

    def append_rows(self, table_index: int, rows_data: list):
        """
        在文档中的第table_index个表格出，添加多行

        :param  table_index:
            需要添加行的表格下标
        :param  rows_data
            需要添加行的数据
        :return:
        """
        for row_data in rows_data:
            self.append_row(table_index, row_data)

    @staticmethod
    def set_font(
            element:
            Document | Table | _Row | _Cell | Paragraph | Run | list[Run],
            font: str, font_size: int,
    ):
        """
        设置元素的字体

        :param element
            需要修改字体字号的元素
        :param font str
            字体
        :param font_size int
            字体大小

        """
        runs = part.get_runs_in_element(element)
        for run in runs:
            run.font.name = font
            run.font.size = Pt(font_size)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font)

    def add_paragraph(self, content: str = None, flag: str = None) \
            -> Paragraph:
        """
        在flag段落前面添加一个段落并返回
        :param flag str
            需要在段落前添加段落的段落的文本
            flag需要为一个单独的段落的文本
            flag建议为全局唯一的文本
            如果flag为None，则在文章末尾添加段落
        :param content str
            需要添加段落的文本
        :return 新添加的段落
        """

        if flag:
            for paragraph in self.paragraphs:
                if flag == paragraph.text:
                    new_p = paragraph.insert_paragraph_before(text=content)
                    self.paragraphs = part.get_paragraphs(self.doc)
                    return new_p

        else:
            new_p = self.doc.add_paragraph(content)
            self.paragraphs = part.get_paragraphs(self.doc)
            return new_p

    @staticmethod
    def indent(paragraph: Paragraph, num: int):
        """
        段落首行缩进多个字符
        :param paragraph docx.text.paragraph.Paragraph
            需要修饰的段落
        :param num int
            需要缩进的字符个数
        """
        paragraph.paragraph_format.first_line_indent = \
            paragraph.style.font.size * num

    def delete_run(self, text: str):
        """
        删除内容text的蚊子块

        :param text: 文字快内容
        :return:
        """
        for run in self.runs:
            if run.text == text:
                run.text = ''
                self.runs.remove(run)
                break

    @staticmethod
    def set_paragraph_align(paragraph: Paragraph, align):
        """
        设置段落的对齐方式

        :param paragraph docx.paragraph.Paragraph
            需要设置对齐方式的段落
        :param align DocDecorate里面的属性
            对齐方式，在DocDecorate类中定义的静态属性
        """
        paragraph.alignment = align

    @staticmethod
    def set_table_align(table: Table, align: int):
        """
        设置表格对齐方式

        :param table docx.table.Table
            需要对齐的表格
        :param align
            对齐方式，DocDecorate类提供的静态属性
            align也可以为数字：
                0: left
                1: center
                2: right
        """
        paragraphs = part.get_paragraphs_in_table(table)
        for paragraph in paragraphs:
            DocDecorate.set_paragraph_align(paragraph, align)

    @staticmethod
    def set_row_align(row: _Row, align: int):
        """
        设置表格里面某一行的对齐方式

        :param row docx.table._Row
            需要设置对齐方式的行
        :param align int
            对齐方式，DocDecorate类里面定义的静态属性
            也可以是数字：
                0: left
                1: center
                2: right
        """
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                DocDecorate.set_paragraph_align(paragraph, align)

    @staticmethod
    def set_column_align(column: _Column, align: int):
        """
        设置表格里面某一列的对齐方式

        :param column docx.table._column
            需要设置对齐方式的列
        :param align int
            对齐方式，DocDecorate类里面定义的静态属性
            也可以是数字：
                0: left
                1: center
                2: right
        """
        for cell in column.cells:
            for paragraph in cell.paragraphs:
                DocDecorate.set_paragraph_align(paragraph, align)

    def add_table(self, rows: int, columns: int) -> Table:
        """
        在文档末尾添加表格

        :param rows int
            表格行数
        :param columns int
            表格列数
        :return docx.table.Table
            返回添加的表格
        """
        table = self.doc.add_table(rows, columns)
        self.tables = part.get_tables(self.doc)
        return table

    @staticmethod
    def add_cell_text(cell: _Cell, text: str) -> _Cell:
        """
        为单元格添加内容

        :param cell: docx.table._Cell
            需要添加文字的单元格对象
        :param text str
            需要添加的内容
        :return _Cell
            单元格对象
        """
        cell.text = text

    @staticmethod
    def add_row_text(row: _Row, content: list[str]) -> _Row:
        """
        为行添加内容

        :param row: docx.table._Row
            需要添加文字的表格行对象
        :param content list[str]
            需要添加的内容列表
        :return _Row
            表格行对象
        """
        assert len(content) == len(row.cells), \
            ValueError('content长度于单元格个数不符')
        for index in range(len(content)):
            row.cells[index].text = str(content[index])
        return row
    
    @staticmethod
    def add_column_text(column: _Column, content: list[str]) -> _Column:
        """
        为列添加内容

        :param column: docx.table._Column
            需要添加文字的表格列对象
        :param content list[str]
            需要添加的内容列表
        :return _column
            表格列对象
        """
        assert len(content) == len(column.cells), \
            ValueError('content长度于单元格个数不符')
        for index in range(len(content)):
            column.cells[index].text = str(content[index])
        return column

    def get_doc(self):
        return self.doc
