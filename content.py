"""
操作文档内容模块：

 - 添加段落
 - 添加表格
 - 添加图片
 - 替换文本
"""
import docx
from docx.document import Document
from docx.shared import Cm
from docx.styles.style import (
    _ParagraphStyle,
    _TableStyle,
)
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from element import DOCXElement


class DocContent(object):
    def __init__(self, doc: str | Document = None):
        if isinstance(doc, str):
            self.doc = docx.Document(doc)
        elif isinstance(doc, DocContent):
            self.doc = doc
        else:
            self.doc = docx.Document()
        self.doc:DocContent

    def add_paragraph(self, text: str, style: _ParagraphStyle = None):
        """
        在文本末尾添加一个段落
        :param text: 段落文本
        :param style: 段落样式，当style为None时默认样式为Normal
        :return: 新添加的段落对象
        """
        paragraph = self.doc.add_paragraph(text, style)
        return paragraph

    def insert_paragraph_before(self, text: str, position: Paragraph | int, style: _ParagraphStyle = None) -> Paragraph:
        """
        在指定段落前面添加一个段落
        :param text: 段落文本内容
        :param position: 指定段落对象或者段落下标
        :param style: 段落样式
        :return: 插入的段落对象
        """
        pos_paragraph = DOCXElement(self.doc).get_paragraph(position)
        paragraph = pos_paragraph.insert_paragraph_before(text, style)
        return paragraph

    def append_text_to_paragraph(self, text: str, position: Paragraph | int, style: _ParagraphStyle = None) -> Run:
        """
        在段落后面添加文本
        :param text: 文本内容
        :param position: 指定段落对象
        :param style: 段落样式
        :return: run对象
        """
        paragraph = DOCXElement(self.doc).get_paragraph(position)
        run = paragraph.add_run(text, style)
        return run

    def add_picture(self, picture: str, width: float = 15, height: float = None):
        """
        在文章末尾添加图片
        :param picture:图片路径
        :param width: 图片宽度，单位厘米，默认15厘米
        :param height: 图片高度，单位厘米。默认None：根据宽度自适应
        :return: InlineShape对象
        """
        self.doc.add_picture(picture, Cm(width), Cm(height) if height else None)

    @staticmethod
    def add_picture_after_paragraph(picture: str, paragraph: Paragraph, width: float = 15,
                                    height: int = None):
        """
        在一个段落的末尾添加图片
        :param picture: 图片路径
        :param paragraph: 段落对象
        :param width: 图片宽度，单位厘米，默认15厘米
        :param height: 图片高度，单位厘米。默认None：根据宽度自适应
        :return: InlineShape对象
        """
        run = paragraph.add_run()
        run.add_picture(picture, Cm(width), Cm(height) if height else None)

    def add_picture_in_run(self, picture: str, run: Run | str, width: float = 15, height: int = None):
        """
        在run对象里添加图片
        :param picture: 图片路径
        :param run: 需要添加图片的run对象或者run里的文本，文本建议为全局唯一的文字
        :param width: 图片宽度，单位厘米，默认15厘米
        :param height: 图片高度，单位厘米。默认None：根据宽度自适应
        :return: InlineShape对象
        """
        if isinstance(run, str):
            runs = DOCXElement(self.doc).get_all_runs()
            for item in runs:
                if item.text == run:
                    run = item
                    break
            else:
                return
        run.add_picture(picture, Cm(width), Cm(height) if height else None)

    @staticmethod
    def add_picture_in_cell(picture: str, cell: _Cell, width: float = 15, height: int = None):
        """
        在单元格中添加图片
        :param picture: 图片路径
        :param cell: 需要添加图片的单元格对象
        :param width: 图片宽度，单位厘米，默认15厘米
        :param height: 图片高度，单位厘米。默认None：根据宽度自适应
        :return: InlineShape对象
        """
        last_paragraph = DOCXElement.get_paragraphs_of_cell(cell)[-1]
        run = last_paragraph.add_run()
        run.add_picture(picture, Cm(width), Cm(height) if height else None)

    def add_table(self, rows: int, columns: int, data: list, style: _TableStyle = None):
        """
        在文章末尾添加表格
        :param rows: 行数
        :param columns: 列数
        :param data: 数据
        :param style: 表格样式
        :return: 表格对象
        """
        assert rows > 0 and columns > 0, ValueError('表格行数或列数不能小于或者等于0')
        assert len(data) >= rows * columns, ValueError('数据个数应该大于单元格个数')
        table = self.doc.add_table(rows, columns, style)
        cells = DOCXElement(self.doc).get_cells_of_table(table)
        for index in range(rows * columns):
            cells[index].text = data[index]
        return table

    @staticmethod
    def append_rows(table: Table, data: list):
        """
        在表格后面添加行
        :param table: 表格对象
        :param data: 数据
        :return: None
        """
        column_length = len(table.rows[0])
        rows = int(len(data) / column_length)
        for row in range(rows):
            row = table.add_row()
            for index in range(column_length):
                row.cells[index].text = data[row * column_length + index]

    @staticmethod
    def clear_paragraph(paragraph: Paragraph):
        """
        清除段落
        :param paragraph: 段落对象或者段落下标
        :return:
        """
        paragraph.clear()
